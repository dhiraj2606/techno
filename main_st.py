import time
from io import BytesIO
from selenium import webdriver
from PIL import Image
from docx import Document
from docx.shared import Inches
from selenium.webdriver.common.by import By
import streamlit as st
import pandas as pd

def process_symbols(symbols):
    # Join the symbols with commas
    joined_symbols = ','.join(symbols)
    # Replace ' -' and '&' with '_'
    processed_symbols = joined_symbols.replace(' -', '_').replace('&', '_')
    return processed_symbols

def create_hyperlink(symbol):
    return f"https://chartink.com/stocks/{symbol}.html"

def take_screenshot_and_save(driver, symbol, url, first_symbol=False, period='daily'):
    driver.get(url)
    time.sleep(2)  # Allow time for the webpage to fully load
    driver.maximize_window()
    time.sleep(1)
    if first_symbol:
        # Click on period
        driver.find_element(By.XPATH, "/html/body/div[2]/div/div[2]/div/div[2]/div[2]/div/div/div[3]/form/div[1]/div[2]/select").click()
        time.sleep(1)
        if period == 'daily':
            driver.find_element(By.XPATH, "/html/body/div[2]/div/div[2]/div/div[2]/div[2]/div/div/div[3]/form/div[1]/div[2]/select/option[11]").click()
        elif period == 'weekly':
            driver.find_element(By.XPATH, "/html/body/div[2]/div/div[2]/div/div[2]/div[2]/div/div/div[3]/form/div[1]/div[2]/select/option[16]").click()
        elif period == 'monthly':
            driver.find_element(By.XPATH, "/html/body/div[2]/div/div[2]/div/div[2]/div[2]/div/div/div[3]/form/div[1]/div[2]/select/option[17]").click()
        time.sleep(1)
        # Click on range
        driver.find_element(By.XPATH, "/html/body/div[2]/div/div[2]/div/div[2]/div[2]/div/div/div[3]/form/div[4]/div/div[2]/select").click()
        time.sleep(1)
        # Select range based on period
        if period == 'daily':
            driver.find_element(By.XPATH, "/html/body/div[2]/div/div[2]/div/div[2]/div[2]/div/div/div[3]/form/div[4]/div/div[2]/select/option[1]").click()
        elif period == 'weekly':
            driver.find_element(By.XPATH, "/html/body/div[2]/div/div[2]/div/div[2]/div[2]/div/div/div[3]/form/div[4]/div/div[2]/select/option[2]").click()
        elif period == 'monthly':
            driver.find_element(By.XPATH, "/html/body/div[2]/div/div[2]/div/div[2]/div[2]/div/div/div[3]/form/div[4]/div/div[2]/select/option[3]").click()
        time.sleep(1)    
        # Click update button
        driver.find_element(By.XPATH, "/html/body/div[2]/div/div[2]/div/div[2]/div[2]/div/div/div[3]/form/input").click()
        # Minimize indicators
        driver.find_element(By.XPATH, "/html/body/div[2]/div/div[1]/div/div[2]/p/u").click()
        time.sleep(1)

    screenshot = driver.get_screenshot_as_png()
    image = Image.open(BytesIO(screenshot))
    return image

def create_word_document(symbols, period):
    driver = webdriver.Chrome()  # Create webdriver instance
    doc = Document()
    first_symbol = True  # Flag to indicate the first symbol
    screenshots = []
    for symbol in symbols:
        url = create_hyperlink(symbol)
        screenshot = take_screenshot_and_save(driver, symbol, url, first_symbol, period)
        screenshots.append((symbol, screenshot, url))
        if first_symbol:
            first_symbol = False  # Update flag for subsequent symbols
        doc.add_paragraph(symbol)
        stream = BytesIO()
        screenshot.save(stream, format="PNG")
        stream.seek(0)
        doc.add_picture(stream, width=Inches(6))
        doc.add_paragraph(url)
    driver.quit()  # Quit webdriver after all symbols are processed
    if period == 'daily':
        doc.save("stock_analysis_daily.docx")
    elif period == 'weekly':
        doc.save("stock_analysis_weekly.docx")
    elif period == 'monthly':
        doc.save("stock_analysis_monthly.docx")
    return screenshots

def stock_analysis_app(period, key_prefix):
    st.title(f"{period.capitalize()} Stock Analysis")
    symbols_input = st.text_area("Enter comma-separated symbol list:", key=f"{key_prefix}_symbols_input")
    if st.button("Generate Report", key=f"{key_prefix}_generate_report"):
        if symbols_input:
            symbols = [symbol.strip() for symbol in symbols_input.split(",")]
            screenshots = create_word_document(symbols, period)
            st.success("Word document created successfully!")
            if period == 'daily':
                with open("stock_analysis_daily.docx", "rb") as file:
                    st.download_button(label="Download Report", data=file, file_name="stock_analysis_daily.docx", key=f"{key_prefix}_download_report")
            elif period == 'weekly':
                with open("stock_analysis_weekly.docx", "rb") as file:
                    st.download_button(label="Download Report", data=file, file_name="stock_analysis_weekly.docx", key=f"{key_prefix}_download_report")
            elif period == 'monthly':
                with open("stock_analysis_monthly.docx", "rb") as file:
                    st.download_button(label="Download Report", data=file, file_name="stock_analysis_monthly.docx", key=f"{key_prefix}_download_report")
            return screenshots
        else:
            st.error("Please enter at least one symbol.")
    return None

def display_screenshots(screenshots):
    if screenshots:
        for symbol, screenshot, url in screenshots:
            st.header(symbol)
            st.image(screenshot, caption=symbol)
            st.write(f"[Chartink Link]({url})")


def main():
    st.title("Chartink Technical Breakouts Analysis")
    tabs = st.tabs(["Symbol Processing","Daily TF", "Weekly TF", "Monthly TF", "Daily Report", "Weekly Report", "Monthly Report"])
    
    with tabs[0]:
        symbols_input = st.text_area("Paste your symbols from Excel (one per line):")
        if st.button("Process"):
        # Split symbols by line break
            symbols_list = symbols_input.split('\n')
        # Process the symbols
            processed_output = process_symbols(symbols_list)

        # Output the processed result
            st.write("Processed Symbols:")
            st.write(processed_output)
    with tabs[1]:
        screenshots_d = stock_analysis_app('daily', 'daily')
    
    with tabs[2]:
        screenshots_w = stock_analysis_app('weekly', 'weekly')
    
    with tabs[3]:
        screenshots_m = stock_analysis_app('monthly', 'monthly')
    
    with tabs[4]:
        if screenshots_d:
            display_screenshots(screenshots_d)
        else:
            st.write("No output available yet. Run the Stock Analysis first.")
    with tabs[5]:
        if screenshots_w:
            display_screenshots(screenshots_w)
        else:
            st.write("No output available yet. Run the Stock Analysis first.")
    
    with tabs[6]:
        if screenshots_m:
            display_screenshots(screenshots_m)
        else:
            st.write("No output available yet. Run the Stock Analysis first.")
if __name__ == "__main__":
    main()
